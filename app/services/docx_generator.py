from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.schemas import ClassifiedOutput, ArticleWithContent
from app.utils.logging import get_logger

logger = get_logger(__name__)


def generate_docx(
    classification: ClassifiedOutput,
    articles: list[ArticleWithContent],
    output_path: Path,
    date_str: str,
) -> Path:
    """Generate the DOCX table of contents matching the required template.

    Template:
    [더벨]
    1. Deal
       A. 경영권 인수 및 매각, 투자 유치
          (1) article title
       B. 투자회수
          (1) article title
       C. 기타
          (1) article title
    2. Industry
       A. E&F 포트폴리오 관련 산업 업계 동향
          - 환경/폐기물
            (1) article title
          - 건설/부동산
            (1) article title
          - 바이오/헬스케어
            (1) article title
       B. 기타 주요 산업 관련 업계 동향
          (1) article title
    3. Fundraising, LP 이슈 및 GP 선정
       (1) article title
    """
    articles_map = {a.info.id: a for a in articles}

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"(딜사이트플러스) Daily News Clipping {date_str}")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # Blank line

    # Fixed indent levels (cm)
    _INDENT_CAT = Cm(0)       # 1. Deal
    _INDENT_SUB = Cm(0.8)     # A. 경영권 인수 및 매각
    _INDENT_ITEM = Cm(1.6)    # - 환경/폐기물
    _INDENT_ART_L1 = Cm(0.8)  # articles under category (Fundraising)
    _INDENT_ART_L2 = Cm(1.6)  # articles under subcategory
    _INDENT_ART_L3 = Cm(2.4)  # articles under sub-item

    def add_articles(article_ids: list[str], indent: Cm):
        if not article_ids:
            return
        count = 0
        for aid in article_ids:
            a = articles_map.get(aid)
            if a:
                count += 1
                title_text = a.info.title.strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = indent
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(f"({count})  {title_text}")
                run.font.size = Pt(10)

    for cat_num, cat in enumerate(classification.categories, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = _INDENT_CAT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{cat_num}. {cat.name}")
        run.bold = True
        run.font.size = Pt(12)

        if cat.subcategories:
            # Subcategories
            for sub_idx, sub in enumerate(cat.subcategories):
                sub_letter = chr(65 + sub_idx)  # A, B, C...

                p = doc.add_paragraph()
                p.paragraph_format.left_indent = _INDENT_SUB
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"{sub_letter}. {sub.name}")
                run.bold = True
                run.font.size = Pt(11)

                if sub.sub_items:
                    # Sub-items (e.g., 환경/폐기물, 건설/부동산, etc.)
                    for sub_item in sub.sub_items:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = _INDENT_ITEM
                        p.paragraph_format.first_line_indent = Cm(0)
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(1)
                        run = p.add_run(f"- {sub_item.name}")
                        run.font.size = Pt(10)

                        if sub_item.articles:
                            add_articles(sub_item.articles, indent=_INDENT_ART_L3)
                if sub.articles:
                    add_articles(sub.articles, indent=_INDENT_ART_L2)
        if cat.articles:
            add_articles(cat.articles, indent=_INDENT_ART_L1)

    doc.save(str(output_path))
    logger.info(f"DOCX generated: {output_path}")
    return output_path
